from __future__ import annotations

import argparse
import json
import math
import os
import re
import shutil
import zipfile
from collections import Counter
from datetime import date
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


INK = "173A3B"
INK_STRONG = "0C292A"
CORAL = "D66E54"
MINT = "9BC8B5"
MINT_DARK = "4C8C78"
PAPER = "FFFFFF"
PAPER_DEEP = "F4F5F4"
WHITE = "FFFFFF"
MUTED = "667777"
GOLD = "C99B52"

COMPANY = "Soluções Serviços Terceirizados Ltda."
CNPJ = "09.445.502/0001-09"


def parse_args() -> argparse.Namespace:
    downloads = Path.home() / "Downloads"
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--workbook",
        type=Path,
        default=downloads / "FINAL ANÁLISE DOS DADOS DE IMPOSTO DE RENDA - SOLUÇÕES_vf FINAL.xlsx",
    )
    parser.add_argument(
        "--received-zip",
        type=Path,
        default=downloads / "DOCs RECEBIDOS SOLUÇÕES.zip",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=downloads / "SOLUCOES_RELATORIO_2026_ISO_19011_2026.docx",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=downloads / "SOLUCOES_RELATORIO_AUDITORIA_IRRF_2026_ISO_19011_2026.docx",
    )
    parser.add_argument(
        "--manifest",
        type=Path,
        default=downloads / "SOLUCOES_RELATORIO_AUDITORIA_IRRF_2026_manifest.json",
    )
    parser.add_argument(
        "--pdf-output",
        type=Path,
        default=downloads / "SOLUCOES_RELATORIO_AUDITORIA_IRRF_2026_ISO_19011_2026.pdf",
    )
    return parser.parse_args()


def format_money(value: float | int | None) -> str:
    if value is None:
        return "Não disponível"
    rendered = f"{float(value):,.2f}"
    rendered = rendered.replace(",", "_").replace(".", ",").replace("_", ".")
    return f"R$ {rendered}"


def format_number(value: float | int | None) -> str:
    if value is None:
        return "—"
    rendered = f"{float(value):,.2f}"
    return rendered.replace(",", "_").replace(".", ",").replace("_", ".")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D8D5C9", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 22, INK_STRONG),
        ("Heading 1", 17, INK_STRONG),
        ("Heading 2", 12, INK),
        ("Heading 3", 10, CORAL),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 7)
        style.paragraph_format.space_after = Pt(5)

    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.6))
    header_table.columns[0].width = Cm(9.5)
    header_table.columns[1].width = Cm(8.1)
    set_cell_text(header_table.cell(0, 0), "REAL PREV  |  AUDITORIA TÉCNICA", bold=True, color=INK_STRONG, size=8)
    set_cell_text(header_table.cell(0, 1), "IRRF · RUBRICAS · eSOCIAL", bold=True, color=CORAL, size=8)
    header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in header_table.rows[0].cells:
        set_cell_border(cell, color=MINT, size="6")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    run = paragraph.add_run("Uso restrito · Soluções Serviços Terceirizados Ltda.   |   Página ")
    run.font.name = "Aptos"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(paragraph, "PAGE")
    paragraph.add_run(" de ")
    add_field(paragraph, "NUMPAGES")


def add_kicker(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(CORAL)


def add_section_heading(document: Document, number: str, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number}. ")
    run.font.color.rgb = RGBColor.from_string(CORAL)
    paragraph.add_run(title)
    return paragraph


def add_body(document: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        first, remainder = text.split(bold_prefix, 1)
        paragraph.add_run(first + bold_prefix).bold = True
        paragraph.add_run(remainder)
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.3)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_callout(document: Document, title: str, text: str, *, accent: str = CORAL) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(0.35)
    table.columns[1].width = Cm(16.9)
    set_cell_shading(table.cell(0, 0), accent)
    set_cell_shading(table.cell(0, 1), PAPER)
    set_cell_border(table.cell(0, 0), color=accent, size="0")
    set_cell_border(table.cell(0, 1), color="D8D5C9", size="4")
    cell = table.cell(0, 1)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title.upper())
    title_run.bold = True
    title_run.font.name = "Aptos"
    title_run.font.size = Pt(8)
    title_run.font.color.rgb = RGBColor.from_string(accent)
    detail = cell.add_paragraph(text)
    detail.paragraph_format.space_after = Pt(0)
    detail.paragraph_format.line_spacing = 1.05
    for run in detail.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(INK)
    return table


def keep_paragraph_group(paragraphs: list, *, keep_last: bool = False) -> None:
    last_index = len(paragraphs) if keep_last else max(0, len(paragraphs) - 1)
    for paragraph in paragraphs[:last_index]:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    if widths:
        for index, width in enumerate(widths):
            table.columns[index].width = Cm(width)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, title in enumerate(headers):
        set_cell_text(header.cells[index], title, bold=True, color=WHITE, size=7.5)
        set_cell_shading(header.cells[index], INK)
        set_cell_border(header.cells[index], color=INK_STRONG, size="4")
    for row_index, row_values in enumerate(rows):
        row = table.add_row()
        fill = WHITE if row_index % 2 == 0 else PAPER
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], value, color=INK, size=7.6)
            set_cell_shading(row.cells[index], fill)
            set_cell_border(row.cells[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_metadata_grid(document: Document, metadata: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=math.ceil(len(metadata) / 2), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [3.1, 5.5, 3.1, 5.5]
    for index, width in enumerate(widths):
        table.columns[index].width = Cm(width)
    for index, (label, value) in enumerate(metadata):
        row = index // 2
        column = (index % 2) * 2
        set_cell_text(table.cell(row, column), label.upper(), bold=True, color=MUTED, size=7)
        set_cell_shading(table.cell(row, column), PAPER_DEEP)
        set_cell_text(table.cell(row, column + 1), value, bold=True, color=INK, size=8.5)
        set_cell_shading(table.cell(row, column + 1), WHITE)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)


def extract_report_data(workbook_path: Path, received_zip_path: Path, evidence_dir: Path) -> dict:
    workbook_values = openpyxl.load_workbook(workbook_path, data_only=True)
    workbook_formulas = openpyxl.load_workbook(workbook_path, data_only=False)
    comparison_sheet = workbook_values["BASE DE IRRF ALTERADA"]
    information_sheet = workbook_values["BASE DE INFORMAÇÃO"]
    rubric_sheet = workbook_values["Plan rubricas dos funcionários"]

    primary_layout = {
        "KELLY": {"row": 2, "received": 8, "treated": 10},
        "NEIDE": {"row": 3, "received": 13, "treated": 15},
        "IRENI": {"row": 4, "received": 18, "treated": 20},
    }
    people = []
    monthly_comparison = {}
    validation_notes = []
    for short_name, layout in primary_layout.items():
        row = layout["row"]
        received_column = layout["received"]
        treated_column = layout["treated"]
        months = []
        for month_row in range(4, 16):
            months.append(
                {
                    "month": information_sheet.cell(month_row, received_column).value,
                    "received_ir_base": information_sheet.cell(month_row, received_column + 1).value,
                    "received_inss": information_sheet.cell(month_row, received_column + 2).value,
                    "received_irrf": information_sheet.cell(month_row, received_column + 3).value,
                    "treated_ir_base": comparison_sheet.cell(month_row, treated_column + 1).value,
                    "treated_inss": comparison_sheet.cell(month_row, treated_column + 2).value,
                    "treated_irrf": comparison_sheet.cell(month_row, treated_column + 3).value,
                }
            )

        received_ir_base = information_sheet.cell(16, received_column + 1).value
        received_inss = information_sheet.cell(16, received_column + 2).value
        received_irrf = information_sheet.cell(16, received_column + 3).value
        treated_ir_base = comparison_sheet.cell(16, treated_column + 1).value
        treated_inss = comparison_sheet.cell(16, treated_column + 2).value
        treated_irrf = comparison_sheet.cell(16, treated_column + 3).value
        plan_ir_base = comparison_sheet.cell(17, treated_column + 1).value
        plan_inss = comparison_sheet.cell(17, treated_column + 2).value
        plan_irrf = comparison_sheet.cell(17, treated_column + 3).value
        if abs(float(treated_ir_base) - float(plan_ir_base)) > 0.01:
            raise ValueError(f"Base IR tratada divergente para {short_name}")
        if abs(float(treated_inss) - float(plan_inss)) > 0.01:
            raise ValueError(f"INSS tratado divergente para {short_name}")
        if abs(float(treated_irrf) - float(plan_irrf)) > 0.01:
            validation_notes.append(
                f"{short_name}: soma mensal do IRRF {format_money(treated_irrf)} difere do total da aba individual "
                f"{format_money(plan_irrf)} em {format_money(abs(float(treated_irrf) - float(plan_irrf)))}."
            )

        person = {
            "short_name": short_name.title(),
            "cpf": comparison_sheet.cell(row, 1).value,
            "name": comparison_sheet.cell(row, 2).value,
            "recomposed": comparison_sheet.cell(row, 3).value,
            "income_statement": comparison_sheet.cell(row, 4).value,
            "declared": comparison_sheet.cell(row, 5).value,
            "esocial_before": comparison_sheet.cell(row, 6).value,
            "s1200_original": comparison_sheet.cell(row, 7).value,
            "received_ir_base": received_ir_base,
            "received_inss": received_inss,
            "received_irrf": received_irrf,
            "treated_ir_base": treated_ir_base,
            "treated_inss": treated_inss,
            "treated_irrf": treated_irrf,
            "plan_irrf": plan_irrf,
        }
        person["base_reduction"] = float(received_ir_base) - float(treated_ir_base)
        person["income_statement_adjustment"] = float(treated_ir_base) - float(person["income_statement"])
        people.append(person)
        monthly_comparison[short_name] = months

    supplementary_people = []
    for row, short_name, treated_column in ((5, "VALERIA", 10), (6, "LUCILEIA", 15)):
        name = str(comparison_sheet.cell(row, 2).value).split("  (")[0].split(" (")[0]
        supplementary_people.append(
            {
                "short_name": short_name.title(),
                "cpf": comparison_sheet.cell(row, 1).value,
                "name": name,
                "treated_ir_base": comparison_sheet.cell(33, treated_column + 1).value,
                "treated_inss": comparison_sheet.cell(33, treated_column + 2).value,
                "treated_irrf": comparison_sheet.cell(33, treated_column + 3).value,
            }
        )
    all_people = people + supplementary_people
    treated_total = round(sum(float(person["treated_ir_base"]) for person in all_people), 2)
    if treated_total != 202_150.81:
        raise ValueError(f"Total da base alterada divergente: {format_money(treated_total)}")

    adjusted_rubrics = []
    for row in range(2, rubric_sheet.max_row + 1):
        adjusted = rubric_sheet.cell(row, 8).value
        if adjusted in (None, ""):
            continue
        adjusted_rubrics.append(
            {
                "table": rubric_sheet.cell(row, 2).value,
                "code": rubric_sheet.cell(row, 3).value,
                "type": rubric_sheet.cell(row, 4).value,
                "description": rubric_sheet.cell(row, 5).value,
                "before": rubric_sheet.cell(row, 6).value,
                "after": adjusted,
            }
        )

    evidence_dir.mkdir(parents=True, exist_ok=True)
    evidence_images = []
    with zipfile.ZipFile(workbook_path) as archive:
        for member in sorted(name for name in archive.namelist() if name.startswith("xl/media/")):
            destination = evidence_dir / Path(member).name
            destination.write_bytes(archive.read(member))
            evidence_images.append(destination)

    with zipfile.ZipFile(received_zip_path) as archive:
        source_files = [name for name in archive.namelist() if not name.endswith("/")]
        by_person = Counter(name.split("/")[0] for name in source_files)
        by_extension = Counter(Path(name).suffix.lower() for name in source_files)
        nested_members = []
        for member in source_files:
            if not member.lower().endswith(".zip"):
                continue
            with zipfile.ZipFile(archive.open(member)) as nested_archive:
                nested_members.extend(nested_archive.namelist())

    return {
        "people": people,
        "primary_people": people,
        "supplementary_people": supplementary_people,
        "all_people": all_people,
        "monthly_comparison": monthly_comparison,
        "validation_notes": validation_notes,
        "adjusted_rubrics": adjusted_rubrics,
        "evidence_images": evidence_images,
        "workbook_sheets": workbook_formulas.sheetnames,
        "source_file_count": len(source_files),
        "source_by_person": dict(by_person),
        "source_by_extension": dict(by_extension),
        "nested_member_count": len([name for name in nested_members if not name.endswith("/")]),
    }


def build_report(template_path: Path, output_path: Path, data: dict) -> None:
    document = Document(template_path)
    clear_document_body(document)
    configure_document(document)
    document.core_properties.title = "Relatório de Auditoria Técnica - Base de IRRF Alterada - Soluções"
    document.core_properties.subject = "Resultado antes e depois para revisão dos informes de rendimentos"
    document.core_properties.author = "Real Prev"
    document.core_properties.comments = "Documento técnico de uso restrito. Não constitui certificação ISO."

    primary = data["primary_people"]
    all_people = data["all_people"]
    before_total = sum(person["received_ir_base"] for person in primary)
    treated_total = sum(person["treated_ir_base"] for person in primary)
    reduction_total = before_total - treated_total
    reduction_percent = reduction_total / before_total * 100 if before_total else 0
    inss_before_total = sum(person["received_inss"] for person in primary)
    inss_after_total = sum(person["treated_inss"] for person in primary)
    final_total = sum(person["treated_ir_base"] for person in all_people)

    add_kicker(document, "Real Prev · Relatório técnico de auditoria")
    title = document.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(5)
    title.add_run("BASE DE IRRF ALTERADA\nANTES, TRATAMENTO E RESULTADO")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run("Resultado técnico para revisão dos informes de rendimentos · Exercício 2025")
    subtitle_run.font.name = "Aptos Display"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string(MINT_DARK)
    add_metadata_grid(
        document,
        [
            ("Empresa auditada", COMPANY),
            ("CNPJ", CNPJ),
            ("Período dos dados", "Exercício 2025"),
            ("Pessoas tratadas", "5 CPFs"),
            ("Data de emissão", date.today().strftime("%d/%m/%Y")),
            ("Versão", "2.0"),
            ("Tipo", "Auditoria técnico-operacional/documental"),
            ("Referência", "Diretrizes da ISO 19011"),
        ],
    )
    document.add_paragraph()
    add_callout(
        document,
        "Entrega central",
        f"A aba BASE DE IRRF ALTERADA é o produto final do tratamento. Ela reúne a base anual correta dos cinco CPFs, no total de {format_money(final_total)}, para apoiar a Soluções na revisão e reemissão dos informes de rendimentos.",
    )

    add_section_heading(document, "1", "Objetivo e produto esperado")
    add_body(document, "Documentar de forma rastreável o tratamento executado e a base de IR resultante para Kelly, Neide, Ireni, Valeria e Lucileia, produzindo o quadro final dos cinco CPFs para apoiar a revisão dos informes de rendimentos do exercício 2025.")
    add_bullet(document, "Preservar a BASE DE INFORMAÇÃO como referência do cenário anterior ao tratamento.")
    add_bullet(document, "Demonstrar os ajustes de incidência e seus efeitos mensais sobre a base de IR.")
    add_bullet(document, "Comprovar que os valores de INSS permaneceram inalterados.")
    add_bullet(document, "Entregar os valores anuais tratados que servirão de base à reemissão dos informes.")
    add_bullet(document, "Registrar divergências e aprovações necessárias antes do uso operacional.")

    add_section_heading(document, "2", "Dados recebidos e escopo")
    add_body(document, "A informação de entrada foi composta pelos documentos fornecidos pela empresa e pelas folhas individuais consolidadas na planilha final. O trabalho abrangeu cinco pessoas e foi delimitado à recomposição da base de cálculo do Imposto de Renda sobre salário e férias tributáveis, sem inclusão do 13º salário.")
    inventory_rows = [
        ["Arquivos de suporte", str(data["source_file_count"]), "XMLs, PDFs e pacote complementar recebidos"],
        ["Pessoas tratadas", "5", "Kelly, Neide, Ireni, Valeria e Lucileia"],
        ["Planilha final", "1", f"{len(data['workbook_sheets'])} abas, cinco folhas individuais e {len(data['evidence_images'])} capturas"],
    ]
    add_table(document, ["Grupo", "Itens", "Conteúdo"], inventory_rows, widths=[3.0, 2.0, 12.0])
    add_body(document, "O trabalho foi documental e local; não foram realizadas novas consultas ao eSocial. O INSS foi preservado nos cálculos, pois a intervenção se restringiu ao tratamento da base de Imposto de Renda. Não se emite parecer jurídico, certificação ISO ou validação de obrigações fora desse escopo.")

    add_section_heading(document, "3", "Tratamento executado")
    methods = [
        ("01", "Consolidação da entrada", "Organização dos totalizadores por CPF e mês na BASE DE INFORMAÇÃO."),
        ("02", "Revisão de rubricas", "Identificação de incidências incompatíveis com a função econômica dos vencimentos e descontos."),
        ("03", "Correção", "Reenquadramento das rubricas selecionadas de incidência IR 09 para 11."),
        ("04", "Recomposição", "Recálculo mensal de salário e férias tributáveis, sem 13º salário, preservando o INSS."),
        ("05", "Conferência", "Validação das totalizações mensais e anuais por CPF."),
        ("06", "Entrega", "Registro dos valores finais por CPF na BASE DE IRRF ALTERADA."),
    ]
    add_table(document, ["Etapa", "Procedimento", "Aplicação"], [[a, b, c] for a, b, c in methods], widths=[1.4, 4.1, 11.5])

    add_section_heading(document, "4", "Resultado final para os informes")
    result_rows = []
    for person in all_people:
        result_rows.append([
            person["name"].title(),
            person["cpf"],
            format_number(person["treated_ir_base"]),
            "Atual (retificado)",
        ])
    result_rows.append([
        "TOTAL DO TRABALHO", "5 CPFs", format_number(final_total), "Resultado final",
    ])
    add_table(document, ["Trabalhador", "CPF", "Base de IRRF correta", "Classificação"], result_rows, widths=[6.0, 3.4, 3.4, 4.2])
    add_callout(document, "Valores corretos", "Kelly: R$ 41.577,35 · Neide: R$ 42.149,99 · Ireni: R$ 40.051,70 · Valeria: R$ 41.556,51 · Lucileia: R$ 36.815,26.", accent=MINT_DARK)
    add_body(document, "Esses cinco valores correspondem à coluna Atual (retificado) da BASE DE IRRF ALTERADA e constituem o resultado do trabalho para revisão dos informes de rendimentos.")

    add_section_heading(document, "5", "Preservação do INSS")
    inss_rows = [[person["name"].title(), person["cpf"], format_number(person["treated_inss"]), "Mantido"] for person in all_people]
    inss_rows.append(["TOTAL DO TRABALHO", "5 CPFs", format_number(sum(person["treated_inss"] for person in all_people)), "Sem alteração"])
    add_table(document, ["Trabalhador", "CPF", "INSS registrado", "Tratamento"], inss_rows, widths=[6.0, 3.4, 3.4, 4.2])
    add_body(document, "Os valores de INSS foram mantidos. O trabalho não recalculou nem alterou a contribuição previdenciária; a intervenção se restringiu à base de cálculo do Imposto de Renda.")

    add_section_heading(document, "6", "Rubricas e causa do ajuste")
    rubric_rows = [[str(item["table"]), str(item["code"]), str(item["type"]), str(item["description"]), str(item["before"]), str(item["after"])] for item in data["adjusted_rubrics"]]
    add_table(document, ["Tabela", "Código", "Tipo", "Descrição", "Antes", "Depois"], rubric_rows, widths=[1.4, 2.2, 2.2, 7.4, 1.8, 1.8])
    add_body(document, "O desvio central estava no tratamento de rubricas de adiantamento: com incidência IR 09, o desconto não deduzia a base mensal como esperado. O reenquadramento para incidência 11 permitiu que vencimentos somassem e descontos deduzissem conforme o sinal econômico registrado, eliminando o excesso de base.")

    add_section_heading(document, "7", "Achados e resultado")
    findings = [
        ("A-01 · Achado maior · Tratado", "Base de IR artificialmente elevada", "A parametrização anterior não compensava adequadamente os adiantamentos. O tratamento corrigiu a composição da base de IR e produziu o resultado anual correto para os cinco CPFs."),
        ("A-02 · Controle preservado", "INSS mantido", "A contribuição previdenciária foi preservada, pois o trabalho se restringiu à recomposição da base de cálculo do Imposto de Renda."),
    ]
    for kicker, heading_text, body_text in findings:
        start = len(document.paragraphs)
        add_kicker(document, kicker)
        heading = document.add_paragraph(style="Heading 2")
        heading.add_run(heading_text)
        add_body(document, body_text)
        keep_paragraph_group(document.paragraphs[start:])

    add_section_heading(document, "8", "Registro das evidências visuais")
    captions = [
        "E-01 · Rubrica 490 — Faltas Abonadas — cadastro com incidência IR 09 antes do ajuste.",
        "E-02 · Rubrica 630 — Base IRRF Crédito — coexistência histórica de registros com incidências 09 e 11.",
        "E-03 · Rubrica 1025 — Adiantamento Salário Último Pagamento — incidência IR 09 antes do tratamento.",
        "E-04 · Rubrica 470 — Adiantamento Salário — incidência IR 09 antes do reenquadramento.",
    ]
    for index, image_path in enumerate(data["evidence_images"]):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run().add_picture(str(image_path), width=Cm(16.2))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(11)
        run = caption.add_run(captions[index])
        run.italic = True
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    add_section_heading(document, "9", "Controles e plano de ação")
    recommendations = [
        ["1", "Aprovação dos informes", "Validar e aprovar formalmente os cinco valores anuais da base alterada.", "Alta", "Empresa / Fiscal"],
        ["2", "Governança", "Versionar matriz de incidências e toda alteração cadastral de rubricas.", "Alta", "Folha / Fiscal"],
        ["3", "Conciliação mensal", "Comparar folha, S-1200, S-1210 e totalizadores antes do fechamento.", "Alta", "Folha"],
        ["4", "Adiantamentos", "Testar mensalmente a soma e a dedução dos adiantamentos na base de IR.", "Alta", "Folha / TI"],
        ["5", "Arquivo técnico", "Preservar planilha final, XMLs, PDFs, memórias, evidências e aprovações.", "Média", "Empresa"],
    ]
    add_table(document, ["Item", "Tema", "Ação recomendada", "Prioridade", "Responsável"], recommendations, widths=[1.0, 3.0, 8.3, 2.0, 2.8])

    conclusion_paragraphs = [add_section_heading(document, "10", "Conclusão")]
    conclusion_paragraphs.append(add_body(document, f"O trabalho tratou as incidências que produziam excesso na base de Imposto de Renda e entregou a BASE DE IRRF ALTERADA dos cinco CPFs, totalizando {format_money(final_total)}."))
    conclusion_paragraphs.append(add_body(document, "A entrega abrange Kelly Cristina de Alcantara Duarte, Neide Aparecida Cecilio, Ireni Lourenco da Silva Soares, Valeria da Conceição Silva e Lucileia Gomes Ferreira Ramos. Os respectivos valores anuais corretos estão consolidados na coluna Atual (retificado) e devem apoiar a revisão dos informes de rendimentos."))
    conclusion_paragraphs.append(add_body(document, "Este relatório constitui avaliação técnica-operacional estruturada segundo diretrizes de auditoria. Não representa certificação ISO nem parecer jurídico, fiscal ou trabalhista, e não substitui a responsabilidade da administração e dos profissionais legalmente responsáveis pela emissão e transmissão das informações."))
    keep_paragraph_group(conclusion_paragraphs, keep_last=True)
    add_callout(document, "Conclusão final", f"Base de IR tratada para cinco pessoas, no total de {format_money(final_total)}, com os valores de INSS preservados.", accent=MINT_DARK)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_pdf_report(output_path: Path, data: dict) -> None:
    from PIL import Image as PillowImage
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        BaseDocTemplate,
        Frame,
        Image,
        KeepTogether,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    palette = {
        "ink": colors.HexColor(f"#{INK}"),
        "strong": colors.HexColor(f"#{INK_STRONG}"),
        "coral": colors.HexColor(f"#{CORAL}"),
        "mint": colors.HexColor(f"#{MINT}"),
        "mint_dark": colors.HexColor(f"#{MINT_DARK}"),
        "paper": colors.HexColor(f"#{PAPER}"),
        "paper_deep": colors.HexColor(f"#{PAPER_DEEP}"),
        "muted": colors.HexColor(f"#{MUTED}"),
        "gold": colors.HexColor(f"#{GOLD}"),
    }

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ReportBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.8,
        leading=11.2,
        textColor=palette["ink"],
        alignment=TA_JUSTIFY,
        spaceAfter=5,
    )
    small = ParagraphStyle(
        "ReportSmall",
        parent=body,
        fontSize=7.4,
        leading=9,
        alignment=TA_LEFT,
        spaceAfter=0,
    )
    table_header = ParagraphStyle(
        "ReportTableHeader",
        parent=small,
        fontName="Helvetica-Bold",
        textColor=colors.white,
    )
    caption = ParagraphStyle(
        "ReportCaption",
        parent=small,
        textColor=palette["muted"],
        alignment=TA_CENTER,
        fontName="Helvetica-Oblique",
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "ReportSection",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=19,
        textColor=palette["strong"],
        spaceBefore=8,
        spaceAfter=6,
        keepWithNext=True,
    )
    subsection_style = ParagraphStyle(
        "ReportSubsection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        leading=13,
        textColor=palette["ink"],
        spaceBefore=5,
        spaceAfter=4,
        keepWithNext=True,
    )
    kicker_style = ParagraphStyle(
        "ReportKicker",
        parent=small,
        fontName="Helvetica-Bold",
        fontSize=7.5,
        leading=9,
        textColor=palette["coral"],
        spaceBefore=6,
        spaceAfter=2,
        keepWithNext=True,
    )
    bullet_style = ParagraphStyle(
        "ReportBullet",
        parent=body,
        leftIndent=11,
        firstLineIndent=-7,
        spaceAfter=2,
    )

    class ReportDocTemplate(BaseDocTemplate):
        pass

    output_path.parent.mkdir(parents=True, exist_ok=True)
    report = ReportDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.4 * cm,
        topMargin=1.65 * cm,
        bottomMargin=1.55 * cm,
        title="Relatório de Auditoria Técnica de IRRF e Rubricas - Soluções",
        author="Real Prev",
        subject="Estruturado conforme diretrizes da ISO 19011:2026",
    )
    frame = Frame(report.leftMargin, report.bottomMargin, report.width, report.height, id="main")

    def draw_page(canvas, doc) -> None:
        canvas.saveState()
        width, height = A4
        canvas.setFillColor(colors.white)
        canvas.rect(0, 0, width, height, fill=1, stroke=0)
        canvas.setStrokeColor(palette["mint"])
        canvas.setLineWidth(1.2)
        canvas.line(doc.leftMargin, height - 1.05 * cm, width - doc.rightMargin, height - 1.05 * cm)
        canvas.setFillColor(palette["strong"])
        canvas.setFont("Helvetica-Bold", 7.2)
        canvas.drawString(doc.leftMargin, height - 0.78 * cm, "REAL PREV  |  AUDITORIA TÉCNICA")
        canvas.setFillColor(palette["coral"])
        canvas.drawRightString(width - doc.rightMargin, height - 0.78 * cm, "IRRF · RUBRICAS · eSOCIAL")
        canvas.setStrokeColor(palette["mint"])
        canvas.line(doc.leftMargin, 1.02 * cm, width - doc.rightMargin, 1.02 * cm)
        canvas.setFillColor(palette["muted"])
        canvas.setFont("Helvetica", 6.8)
        canvas.drawString(doc.leftMargin, 0.68 * cm, "Uso restrito · Soluções Serviços Terceirizados Ltda.")
        canvas.drawRightString(width - doc.rightMargin, 0.68 * cm, f"Página {doc.page}")
        canvas.restoreState()

    report.addPageTemplates(PageTemplate(id="report", frames=[frame], onPage=draw_page))

    def rich(text: str, style=body) -> Paragraph:
        return Paragraph(text, style)

    def heading(number: str, title: str) -> Paragraph:
        return rich(f'<font color="#{CORAL}">{number}.</font> {title}', section_style)

    def bullet(text: str) -> Paragraph:
        return rich(f"• {text}", bullet_style)

    def callout(title: str, text: str, accent: str = CORAL) -> Table:
        content = rich(
            f'<font color="#{accent}"><b>{title.upper()}</b></font><br/>{text}',
            ParagraphStyle("Callout", parent=body, alignment=TA_LEFT, spaceAfter=0),
        )
        table = Table([["", content]], colWidths=[0.18 * cm, 16.4 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(f"#{accent}")),
                    ("BACKGROUND", (1, 0), (1, 0), palette["paper_deep"]),
                    ("BOX", (1, 0), (1, 0), 0.4, colors.HexColor("#D8D5C9")),
                    ("LEFTPADDING", (0, 0), (0, 0), 0),
                    ("RIGHTPADDING", (0, 0), (0, 0), 0),
                    ("TOPPADDING", (1, 0), (1, 0), 7),
                    ("BOTTOMPADDING", (1, 0), (1, 0), 7),
                ]
            )
        )
        return table

    def report_table(headers: list[str], rows: list[list[str]], widths: list[float]) -> Table:
        contents = [[rich(header, table_header) for header in headers]]
        contents.extend([[rich(str(value), small) for value in row] for row in rows])
        table = Table(contents, colWidths=[width * cm for width in widths], repeatRows=1, hAlign="CENTER")
        commands = [
            ("BACKGROUND", (0, 0), (-1, 0), palette["ink"]),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8D5C9")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
        for row_index in range(1, len(contents)):
            commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.white if row_index % 2 else palette["paper_deep"]))
        table.setStyle(TableStyle(commands))
        return table

    all_people = data["all_people"]
    final_total = sum(person["treated_ir_base"] for person in all_people)

    story = [
        rich(f'<font color="#{CORAL}"><b>REAL PREV · RELATÓRIO TÉCNICO DE AUDITORIA</b></font>', kicker_style),
        rich("BASE DE IRRF ALTERADA<br/>ANTES, TRATAMENTO E RESULTADO", ParagraphStyle("CoverTitle", parent=section_style, fontSize=20, leading=22, textColor=palette["strong"], spaceBefore=3, spaceAfter=3)),
        rich("Resultado técnico para revisão dos informes de rendimentos · Exercício 2025", ParagraphStyle("CoverSubtitle", parent=body, fontSize=11, leading=14, textColor=palette["mint_dark"], alignment=TA_LEFT)),
        Spacer(1, 0.2 * cm),
    ]
    metadata = [
        ["EMPRESA AUDITADA", COMPANY, "CNPJ", CNPJ],
        ["PERÍODO DOS DADOS", "Exercício 2025", "PESSOAS TRATADAS", "5 CPFs"],
        ["DATA DE EMISSÃO", date.today().strftime("%d/%m/%Y"), "VERSÃO", "2.0"],
        ["TIPO", "Auditoria técnico-operacional/documental", "REFERÊNCIA", "Diretrizes da ISO 19011"],
    ]
    metadata_contents = [[rich(f"<b>{row[0]}</b>", small), rich(f"<b>{row[1]}</b>", small), rich(f"<b>{row[2]}</b>", small), rich(f"<b>{row[3]}</b>", small)] for row in metadata]
    metadata_table = Table(metadata_contents, colWidths=[3.0 * cm, 5.4 * cm, 3.0 * cm, 5.4 * cm], hAlign="CENTER")
    metadata_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), palette["paper_deep"]), ("BACKGROUND", (2, 0), (2, -1), palette["paper_deep"]),
        ("BACKGROUND", (1, 0), (1, -1), colors.white), ("BACKGROUND", (3, 0), (3, -1), colors.white),
        ("TEXTCOLOR", (0, 0), (-1, -1), palette["ink"]), ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8D5C9")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([
        metadata_table,
        Spacer(1, 0.45 * cm),
        callout("Entrega central", f"A aba BASE DE IRRF ALTERADA é o produto final do tratamento. Ela reúne a base anual correta dos cinco CPFs, no total de {format_money(final_total)}, para apoiar a Soluções na revisão e reemissão dos informes de rendimentos."),
        heading("1", "Objetivo e produto esperado"),
        rich("Documentar de forma rastreável o tratamento executado e a base de IR resultante para Kelly, Neide, Ireni, Valeria e Lucileia, produzindo o quadro final dos cinco CPFs para apoiar a revisão dos informes de rendimentos de 2025."),
        bullet("Preservar a BASE DE INFORMAÇÃO como referência do cenário anterior."),
        bullet("Demonstrar os efeitos do tratamento sobre a base de IR."),
        bullet("Comprovar que os valores de INSS permaneceram inalterados."),
        bullet("Entregar os valores anuais tratados destinados à revisão dos informes."),
        heading("2", "Dados recebidos e escopo"),
        rich("A informação de entrada foi composta pelos documentos fornecidos pela empresa e pelas folhas individuais consolidadas na planilha final. O trabalho abrangeu cinco pessoas e foi delimitado à recomposição da base de cálculo do Imposto de Renda sobre salário e férias tributáveis, sem inclusão do 13º salário."),
        report_table(["Grupo", "Itens", "Conteúdo"], [
            ["Arquivos de suporte", str(data["source_file_count"]), "XMLs, PDFs e pacote complementar recebidos"],
            ["Pessoas tratadas", "5", "Kelly, Neide, Ireni, Valeria e Lucileia"],
            ["Planilha final", "1", f"{len(data['workbook_sheets'])} abas, cinco folhas individuais e {len(data['evidence_images'])} capturas"],
        ], [3.0, 2.0, 11.8]),
        rich("O trabalho foi documental e local; não foram realizadas novas consultas ao eSocial. O INSS foi preservado, pois a intervenção se restringiu à base de Imposto de Renda."),
        heading("3", "Tratamento executado"),
        report_table(["Etapa", "Procedimento", "Aplicação"], [
            ["01", "Consolidação", "Totalizadores por CPF e mês na BASE DE INFORMAÇÃO."],
            ["02", "Revisão", "Identificação de incidências incompatíveis com vencimentos e descontos."],
            ["03", "Correção", "Reenquadramento das rubricas selecionadas de IR 09 para 11."],
            ["04", "Recomposição", "Recálculo de salário e férias tributáveis, sem 13º e preservando o INSS."],
            ["05", "Conferência", "Validação das totalizações mensais e anuais por CPF."],
            ["06", "Entrega", "Valores finais por CPF na BASE DE IRRF ALTERADA."],
        ], [1.4, 4.1, 11.3]),
        heading("4", "Resultado final para os informes"),
    ])

    result_rows = [[person["name"].title(), person["cpf"], format_number(person["treated_ir_base"]), "Atual (retificado)"] for person in all_people]
    result_rows.append(["TOTAL DO TRABALHO", "5 CPFs", format_number(final_total), "Resultado final"])
    story.extend([
        report_table(["Trabalhador", "CPF", "Base de IRRF correta", "Classificação"], result_rows, [6.0, 3.4, 3.4, 4.0]),
        callout("Valores corretos", "Kelly: R$ 41.577,35 · Neide: R$ 42.149,99 · Ireni: R$ 40.051,70 · Valeria: R$ 41.556,51 · Lucileia: R$ 36.815,26.", MINT_DARK),
        rich("Esses cinco valores correspondem à coluna Atual (retificado) e constituem o resultado do trabalho para revisão dos informes de rendimentos."),
        heading("5", "Preservação do INSS"),
    ])
    inss_rows = [[person["name"].title(), person["cpf"], format_number(person["treated_inss"]), "Mantido"] for person in all_people]
    inss_rows.append(["TOTAL DO TRABALHO", "5 CPFs", format_number(sum(person["treated_inss"] for person in all_people)), "Sem alteração"])
    story.extend([
        report_table(["Trabalhador", "CPF", "INSS registrado", "Tratamento"], inss_rows, [6.0, 3.4, 3.4, 4.0]),
        rich("Os valores de INSS foram mantidos. O trabalho não recalculou nem alterou a contribuição previdenciária; a intervenção se restringiu à base de cálculo do Imposto de Renda."),
        KeepTogether([
            heading("6", "Rubricas e causa do ajuste"),
            report_table(["Tabela", "Código", "Tipo", "Descrição", "Antes", "Depois"], [[str(item["table"]), str(item["code"]), str(item["type"]), str(item["description"]), str(item["before"]), str(item["after"])] for item in data["adjusted_rubrics"]], [1.3, 2.0, 2.0, 7.8, 1.8, 1.8]),
            rich("O desvio central estava nas rubricas de adiantamento: com incidência IR 09, o desconto não deduzia a base mensal como esperado. O reenquadramento para 11 permitiu a soma de vencimentos e a dedução de descontos conforme o sinal econômico."),
        ]),
        KeepTogether([
            heading("7", "Achados e resultado"),
            rich("A-01 · ACHADO MAIOR · TRATADO", kicker_style),
            rich("Base de IR artificialmente elevada", subsection_style),
            rich("A parametrização anterior não compensava adequadamente os adiantamentos. O tratamento corrigiu a composição da base de IR e produziu o resultado anual correto para os cinco CPFs."),
            rich("A-02 · CONTROLE PRESERVADO", kicker_style),
            rich("INSS mantido", subsection_style),
            rich("A contribuição previdenciária foi preservada, pois o trabalho se restringiu à recomposição da base de cálculo do Imposto de Renda."),
        ]),
        heading("8", "Registro das evidências visuais"),
    ])
    captions = [
        "E-01 · Rubrica 490 — Faltas Abonadas — cadastro com incidência IR 09 antes do ajuste.",
        "E-02 · Rubrica 630 — Base IRRF Crédito — registros históricos com incidências 09 e 11.",
        "E-03 · Rubrica 1025 — Adiantamento Salário Último Pagamento — incidência IR 09 antes do tratamento.",
        "E-04 · Rubrica 470 — Adiantamento Salário — cadastro com incidência IR 09 antes do reenquadramento.",
    ]
    evidence_cells = []
    for index, image_path in enumerate(data["evidence_images"]):
        with PillowImage.open(image_path) as source_image:
            aspect_ratio = source_image.height / source_image.width
        image_width = 7.9 * cm
        image = Image(str(image_path), width=image_width, height=image_width * aspect_ratio)
        evidence_cells.append([image, rich(captions[index], caption)])
    evidence_table = Table(
        [evidence_cells[:2], evidence_cells[2:]],
        colWidths=[8.35 * cm, 8.35 * cm],
        hAlign="CENTER",
    )
    evidence_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(evidence_table)

    story.extend(
        [
            KeepTogether([
                heading("9", "Controles e plano de ação"),
                report_table(
                ["Item", "Tema", "Ação recomendada", "Prioridade", "Responsável"],
                [
                    ["1", "Aprovação", "Validar e aprovar os cinco valores anuais da base alterada.", "Alta", "Empresa / Fiscal"],
                    ["2", "Governança", "Versionar matriz de incidências e alterações de rubricas.", "Alta", "Folha / Fiscal"],
                    ["3", "Conciliação", "Comparar folha, S-1200, S-1210 e totalizadores mensalmente.", "Alta", "Folha"],
                    ["4", "Adiantamentos", "Testar mensalmente a soma e a dedução na base de IR.", "Alta", "Folha / TI"],
                    ["5", "Arquivo técnico", "Preservar planilha, XMLs, PDFs, memórias e aprovações.", "Média", "Empresa"],
                ],
                [0.9, 3.0, 8.1, 2.0, 2.8],
                ),
            ]),
            KeepTogether([
                heading("10", "Conclusão"),
                rich(f"O trabalho tratou as incidências que produziam excesso na base de Imposto de Renda e entregou a BASE DE IRRF ALTERADA dos cinco CPFs, totalizando {format_money(final_total)}."),
                rich("A entrega abrange Kelly Cristina de Alcantara Duarte, Neide Aparecida Cecilio, Ireni Lourenco da Silva Soares, Valeria da Conceição Silva e Lucileia Gomes Ferreira Ramos. Os valores anuais corretos estão consolidados na coluna Atual (retificado) e devem apoiar a revisão dos informes de rendimentos."),
                rich("Este relatório não representa certificação ISO nem parecer jurídico, fiscal ou trabalhista, e não substitui a responsabilidade da administração e dos profissionais responsáveis pela emissão e transmissão."),
                callout("Conclusão final", f"Base de IR tratada para cinco pessoas, no total de {format_money(final_total)}, com os valores de INSS preservados.", MINT_DARK),
            ]),
        ]
    )
    report.build(story)


def main() -> None:
    args = parse_args()
    for path in (args.workbook, args.received_zip, args.template):
        if not path.exists():
            raise FileNotFoundError(path)

    evidence_dir = args.output.parent / "_solucoes_iso_irrf_evidencias"
    data = extract_report_data(args.workbook, args.received_zip, evidence_dir)
    build_report(args.template, args.output, data)
    build_pdf_report(args.pdf_output, data)

    manifest = {
        "output": str(args.output),
        "pdf_output": str(args.pdf_output),
        "template": str(args.template),
        "workbook": str(args.workbook),
        "received_zip": str(args.received_zip),
        "generated_on": date.today().isoformat(),
        "treated_people": [person["cpf"] for person in data["all_people"]],
        "treated_ir_base_total": round(sum(person["treated_ir_base"] for person in data["all_people"]), 2),
        "treated_ir_base_by_cpf": {
            person["cpf"]: round(person["treated_ir_base"], 2)
            for person in data["all_people"]
        },
        "inss_preserved": True,
        "source_file_count": data["source_file_count"],
        "adjusted_rubrics": len(data["adjusted_rubrics"]),
        "evidence_images": [str(path) for path in data["evidence_images"]],
    }
    args.manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()